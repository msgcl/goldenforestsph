from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Golden_Forests_Website_Structure_Presentation_Script.docx"
LOGO = ROOT / "attached_assets" / "logo.png"

GREEN = "17392E"
DEEP_GREEN = "0E241D"
GOLD = "C8A070"
PALE_GOLD = "F5EADB"
CREAM = "FBFCF7"
CHARCOAL = "1B1B1B"
MUTED = "5E6B65"
LIGHT_GREEN = "EAF0EC"
LIGHT_GOLD = "F5EBDD"
WHITE = "FFFFFF"


DOCS = {
    "company": "6. GF Company Overview-July 2026.pdf",
    "agarwood": "6. GF Agarwood Exposé July 2026.pdf",
    "mango": "6. GF Mango Exposé July 2026.pdf",
    "presentation": "6. GF Professional Client Investment presentation - July 2026.pdf",
    "faq": "6. GF-FAQ-July 2026.pdf",
    "ppm": "9. Private Placement Memorandum-Professional Investors-Draft-July 2026.pdf",
}


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_label_paragraph(doc, label, text, style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(label + " ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(GREEN)
    p.add_run(text)
    return p


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.add_run(item)


def add_script_box(doc, paragraphs):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.55)
    shade(cell, LIGHT_GOLD)
    set_cell_margins(cell, 180, 220, 180, 220)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    title = cell.paragraphs[0]
    title.style = doc.styles["Script Label"]
    title.add_run("PRESENTATION SCRIPT — SAY THIS")
    for text in paragraphs:
        p = cell.add_paragraph(style="Script Text")
        p.add_run(text)
    doc.add_paragraph()


def add_source_box(doc, source_lines, note=None):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    shade(cell, LIGHT_GREEN)
    set_cell_margins(cell, 150, 190, 150, 190)
    p = cell.paragraphs[0]
    p.style = doc.styles["Source Text"]
    r = p.add_run("SOURCE LINEAGE — NEW FILES\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(GREEN)
    for i, source in enumerate(source_lines):
        p.add_run(("• " if i == 0 else "\n• ") + source)
    if note:
        rr = p.add_run("\nNote: " + note)
        rr.italic = True
    doc.add_paragraph()


def add_page_section(doc, number, page):
    doc.add_page_break()
    eyebrow = doc.add_paragraph()
    eyebrow.style = doc.styles["Eyebrow"]
    eyebrow.add_run(f"{page['group'].upper()}  /  PAGE {number:02d}")
    h = doc.add_heading(page["title"], level=1)
    if page.get("route"):
        route = doc.add_paragraph()
        route.style = doc.styles["Route"]
        route.add_run(f"Route: {page['route']}")

    add_label_paragraph(doc, "Why this page is here:", page["why"])

    doc.add_heading("What the audience sees", level=2)
    add_bullets(doc, page["inside"])

    if page.get("demo"):
        add_label_paragraph(doc, "Demonstration cue:", page["demo"])

    add_script_box(doc, page["script"])

    add_label_paragraph(doc, "Transition:", page["transition"])
    add_source_box(doc, page["sources"], page.get("source_note"))


pages = [
    {
        "group": "Overview",
        "title": "Home — The Executive Overview",
        "route": "/",
        "why": "This is the orientation page. It establishes CADI as the Philippine operating company, summarizes the dual-crop programme, shows current nursery information, and gives visitors direct paths into the operational detail.",
        "inside": [
            "Hero statement: professionally managed agroforestry investments operated in the Philippines by CADI for Golden Forests’ ring-fenced agarwood and mango sub-funds.",
            "Primary calls to action for the plantation timeline and nursery dashboard.",
            "Featured operations video showing the nursery, plantation development, and operating environment.",
            "Four programme markers: 10-year agarwood lifecycle, mango fruiting from Year 5, July 2026 out-planting, and a 31,000-tree launch inventory.",
            "Live nursery metrics, share-availability cards, deployment summaries, and inventory date. These values come from the website database and may change.",
            "Operations snapshot cards for nursery, plantation, and technology; governance cards for compliance, precision operations, and transparent reporting.",
            "Zambales access and client-visit summary.",
        ],
        "demo": "Begin with the hero, scroll through the four programme markers, pause on the live metrics, then use the operations snapshot to explain how the rest of the website is organised.",
        "script": [
            "“This is the Golden Forests plantation management portal operated by Crassna Agroforestry Development Inc., or CADI, in the Philippines. The homepage gives us the entire programme in one view: a ten-year agarwood cycle, mango production beginning from Year 5, and the July 2026 field-deployment milestone.”",
            "“The launch programme is diversified across 23,000 agarwood trees and 8,000 mango trees. The live cards below are operational rather than promotional: they show the current nursery position, share availability, and the date of the latest inventory update.”",
            "“From here, I can move into the nursery, plantation, technology, governance, and client-visibility sections. This structure lets a visitor begin with the high-level story and then verify each part of the operating model.”",
        ],
        "transition": "“Let me first explain who sits behind the operating platform and how the group structure works.”",
        "sources": [
            f"{DOCS['company']}, p. 1 — group overview, dual-crop assets, technology, governance, and stewardship.",
            f"{DOCS['presentation']}, pp. 5–9, 24–32 — company, crop strategy, structure, technology, sustainability, visibility, and approvals.",
            f"{DOCS['ppm']}, pp. 1–18 — VCC strategy, operating structure, inventory base, governance, and reporting.",
        ],
        "source_note": "Live stock, availability, and inventory-date values are database-driven website data, not fixed figures copied from a PDF.",
    },
    {
        "group": "Overview",
        "title": "About Us — Corporate Identity and Operating Mandate",
        "route": "/about",
        "why": "This page explains the relationship between Golden Forests, the Singapore VCC sub-funds, licensed placement intermediaries, and CADI. It prevents confusion between fund governance and Philippine plantation operations.",
        "inside": [
            "Corporate overview of Golden Forests’ managed Philippine agroforestry strategy.",
            "Explanation that shares are valued on a tree-equivalent basis for allocation and accounting; shareholders do not own a particular tree, planting block, land parcel, or plantation asset.",
            "CADI’s responsibilities: propagation, out-planting, maintenance, inoculation scheduling, harvest support, product sales, and operational reporting.",
            "Four-part operating mandate: dual-crop delivery, regulated access, Agroforestry Intelligence, and one-for-one native-tree stewardship.",
        ],
        "demo": "Use the three corporate-overview paragraphs to separate the fund layer from the operating-company layer, then finish with the four mandate points.",
        "script": [
            "“This page explains the corporate structure clearly. Eligible professional investors subscribe for shares in dedicated, ring-fenced sub-funds of a Singapore Variable Capital Company through appropriately licensed private-placement intermediaries.”",
            "“A share is measured on a tree-equivalent basis for economic allocation and accounting, but it does not give direct ownership of a particular tree, land, or planting block.”",
            "“CADI is the Philippine operating company. Its role covers the physical lifecycle—from nursery propagation and field planting to maintenance, agarwood inoculation scheduling, harvest support, product sales, and operational reporting. The fund manager and administrator remain responsible for fund-level governance, records, and controlled disbursements.”",
        ],
        "transition": "“With the structure established, I’ll introduce the people responsible for governance and execution.”",
        "sources": [
            f"{DOCS['company']}, p. 1 — company structure, vision, mission, differentiators, and share-class explanation.",
            f"{DOCS['ppm']}, pp. 1–4 and 11–17 — VCC strategy, entity roles, controls, and reporting.",
            f"{DOCS['faq']}, pp. 2–6 — company identity, mission, values, sustainability, and ownership explanation.",
        ],
    },
    {
        "group": "Overview",
        "title": "Management — Leadership and Execution Capability",
        "route": "/management",
        "why": "Agroforestry is a long-duration operating programme. This page shows the multidisciplinary leadership behind corporate governance, finance, agronomy, plantation execution, compliance, communications, and administration.",
        "inside": [
            "Management statement describing more than 80 years of combined relevant experience.",
            "Profile groups for Executive Management, Board of Directors, and Senior Management.",
            "Clickable profile cards opening each person’s experience and expertise.",
            "Current seeded directory includes Charles McKenzie, Mark LM Quinn, Cord Kabus-Duprée, Angie Brion, Marciano Gecolea, R.A.G. Ferdinand Domingo, Adele Frances, Romina Dalit, Billy Medel, Kyla Brion, and Mara Sofia Gecolea.",
        ],
        "demo": "Open one executive profile and one agronomy or plantation profile to show that the page provides more than names and titles.",
        "script": [
            "“A long-term biological-asset strategy depends on people who can execute across several disciplines. The management page therefore separates executive leadership, board oversight, and senior operating management.”",
            "“The profiles combine institutional finance and governance with forestry, mango science, nursery and plantation management, regulatory work, communications, and administration. Each card opens into a fuller biography so a visitor can review the experience relevant to the programme.”",
            "“This page is important because it connects every operational claim elsewhere on the website to a named area of responsibility.”",
        ],
        "transition": "“Now I’ll move from the people to the first physical stage of the programme: the nursery.”",
        "sources": [
            f"{DOCS['ppm']}, p. 10 — leadership and execution capability.",
            f"{DOCS['agarwood']}, p. 15 — management experience and scientific partnerships.",
            f"{DOCS['mango']}, pp. 13–14 — management experience and university partnerships.",
            f"{DOCS['presentation']}, p. 25 — professional operations and lifecycle management.",
        ],
        "source_note": "Individual biographies are maintained in the website database/seed data. The PDFs support the team-capability narrative but do not contain every current website biography.",
    },
    {
        "group": "Operations",
        "title": "Nursery Operations — Propagation and Readiness",
        "route": "/nursery",
        "why": "The nursery is where programme readiness becomes visible. This page tracks planting stock, replacement buffers, propagation methods, growth measurements, mortality, and nursery controls before field deployment.",
        "inside": [
            "Live counts for Aquilaria crassna agarwood and Sweet Elena Carabao mango seedlings.",
            "Current share-availability view for Sub-Fund A and Sub-Fund B allocations.",
            "Seedling gallery with photographs and video, managed through the admin asset library.",
            "Propagation explanation: standardized agarwood protocols; grafting and dwarfing, flowering support, pest and disease control for mango.",
            "Twenty-percent surplus planting buffers for mortality and covered replacement events.",
            "Growth dashboard for average agarwood height, mango height, mortality rate, and latest update date.",
            "Nursery technology protocol: smart irrigation, climate control, pest management, and soil analytics.",
        ],
        "demo": "Point out the latest-updated date before quoting any live figure. Open a nursery image or video, then show the growth dashboard and the technology protocol.",
        "script": [
            "“The nursery page is the operational starting point. It shows the planting stock being prepared for field deployment and distinguishes the agarwood and Sweet Elena mango inventories.”",
            "“The figures and media are live website records, so I always refer to the displayed update date. Below that, the page explains how the crops are propagated and why the programme maintains a twenty-percent surplus buffer for mortality and replacement controls.”",
            "“The growth dashboard tracks average seedling height and mortality, while the technology section explains the moisture, climate, pest, and soil controls used before the trees are transferred to the field.”",
        ],
        "transition": "“Once planting stock meets the required condition, the programme moves into the Zambales plantation rollout.”",
        "sources": [
            f"{DOCS['ppm']}, p. 9 — nursery and launch-inventory base.",
            f"{DOCS['agarwood']}, pp. 8–9, 13–14, and 20 — species, spacing, lifecycle management, technology, and replacement controls.",
            f"{DOCS['mango']}, pp. 4–6, 9, 12, and 18 — Sweet Elena propagation, density, lifecycle management, technology, and risk controls.",
            f"{DOCS['faq']}, pp. 8–10 — sustainability, insurance/security, and replacement stock.",
        ],
        "source_note": "The current counts, measurements, mortality rate, and media are operational database records rather than static PDF content.",
    },
    {
        "group": "Operations",
        "title": "Plantation Operations — Zambales Field Execution",
        "route": "/plantation",
        "why": "This page shows how nursery stock becomes a managed plantation. It documents secured land and preparation, spacing, irrigation and monitoring readiness, cassava intercropping, and both crop lifecycles.",
        "inside": [
            "Field gallery showing secured land, preparation, and deployment readiness.",
            "July 2026 deployment milestone for 23,000 agarwood and 8,000 mango trees, with 4,600 and 1,600 replacement trees respectively.",
            "Spacing: agarwood at 3 m × 2 m, approximately 1,667 trees per hectare; mango at 6 m × 4 m, approximately 416 trees per hectare.",
            "Land preparation using deep ripping, levelling, contour drainage, and terracing.",
            "Cassava intercropping during Years 1–2 to shade young agarwood.",
            "Agarwood lifecycle: establishment, monitoring and maintenance, controlled inoculation, then harvest and extraction in Years 9–10.",
            "Mango lifecycle: grafting and nursery, establishment, fruiting ramp-up from Year 5, then stable annual production through Year 25.",
        ],
        "demo": "Use the milestone cards to explain scale and spacing, then compare the two lifecycle columns to show how the dual-crop timelines complement one another.",
        "script": [
            "“This page follows the stock from nursery to field execution in Zambales. The July 2026 launch inventory is 31,000 production trees: 23,000 Aquilaria crassna and 8,000 Sweet Elena mango, supported by separate twenty-percent replacement buffers.”",
            "“The field design is crop-specific. Agarwood uses three-by-two-metre spacing with cassava between rows during the first two years. Mango uses six-by-four-metre spacing for a high-density but manageable dwarf orchard.”",
            "“The lifecycle comparison is central to the strategy. Agarwood is a ten-year resin cycle with inoculation in Years 7–8 and realization in Years 9–10. Mango develops earlier recurring harvest income, beginning from Year 5 and continuing through the long orchard life.”",
        ],
        "transition": "“Field execution must sit inside a documented legal and regulatory framework, which is the next page.”",
        "sources": [
            f"{DOCS['agarwood']}, pp. 8–9 and 13–14 — plantation design, spacing, cassava, management, and technology.",
            f"{DOCS['mango']}, pp. 4–5 and 9–12 — dwarf orchard density, lifecycle, harvest profile, and technology.",
            f"{DOCS['presentation']}, pp. 12–13 and 19 — agarwood and mango plantation models.",
            f"{DOCS['ppm']}, pp. 9 and 13–16 — inventory, operating responsibilities, controls, and reserve-funded later stages.",
        ],
    },
    {
        "group": "Operations",
        "title": "Compliance Framework — Permits, Trade, and Governance",
        "route": "/compliance",
        "why": "Agarwood cultivation and trade are regulated, and professional-client structures require documented governance. This page gathers permitting, CITES-aligned traceability, supply-chain certification pathways, and client controls in one place.",
        "inside": [
            "Permitting and regulatory-position overview.",
            "Highlights covering Aquilaria crassna permitting foundations, CITES-aligned export planning, ring-fenced fund controls, PEFC/FSC pathways, wildlife-culture permitting, and Bureau of Customs certification.",
            "Four compliance pillars: DENR approval, CITES compliance, supply-chain accreditation, and client-governance controls.",
            "Further mandates for targeted PEFC mango certification, FSC agarwood certification, and definitive contract/governance controls.",
        ],
        "demo": "Explain the distinction between operating permits, export/traceability requirements, certification targets, and fund-governance documents.",
        "script": [
            "“Compliance is presented as an operating system, not a single certificate. At plantation level, CADI works through Philippine environmental, wildlife, customs, and agricultural requirements. For agarwood trade, the workflow also has to support CITES-aligned documentation and traceability.”",
            "“The website separately identifies supply-chain certification targets, including the PEFC and FSC pathways. These are stated as targets rather than completed certifications.”",
            "“At client level, the controls include ring-fenced sub-fund accounting, shareholder records, licensed fund oversight, definitive subscription documents, periodic reporting, and independent audit processes.”",
        ],
        "transition": "“The evidence supporting those operational and reporting controls is strengthened by the programme’s technology stack.”",
        "sources": [
            f"{DOCS['company']}, pp. 1–2 — DENR, CITES, scientific partnerships, and disclosure context.",
            f"{DOCS['agarwood']}, pp. 7, 12–16, and 20 — CITES constraints, licensed technology, traceability, sustainability, and risk management.",
            f"{DOCS['presentation']}, pp. 24, 31–32 — sub-fund structure, client visibility, and regulatory approvals.",
            f"{DOCS['ppm']}, pp. 11–21 and 26–29 — entity controls, reserve governance, reporting, selling restrictions, risks, and launch conditions.",
        ],
    },
    {
        "group": "Operations",
        "title": "Agroforestry Intelligence — AI Precision Farming",
        "route": "/technology",
        "why": "This page explains how field data becomes earlier intervention, more efficient input use, measurable evidence, and scalable operating standards across two different crop programmes.",
        "inside": [
            "Technology stack: drones and yield optimisation, IoT soil sensors, smart irrigation, and AI-enabled monitoring.",
            "Crop-health mapping and targeted interventions.",
            "Moisture and nutrient data at field-block level.",
            "Sensor-led irrigation and input control.",
            "Anomaly detection and predictive alerts for stress, disease, and operational risk.",
            "Operational benefits including better decisions, reduced field-loss risk, efficient inputs, stronger evidence, and scalable standards.",
        ],
        "demo": "Describe the data flow from sensor or drone observation, to alert, to field intervention, to evidence in operational reporting.",
        "script": [
            "“In this website, AI means Agroforestry Intelligence: a practical framework combining agronomic science, sensors, drones, and monitoring software.”",
            "“Drones provide crop-health mapping, soil sensors provide moisture and nutrient visibility, smart irrigation controls water use, and the monitoring layer flags anomalies early.”",
            "“The purpose is not technology for its own sake. It is to help field teams intervene earlier, document what happened at planting-block level, use inputs more efficiently, and create stronger evidence for operational and fund reporting.”",
        ],
        "transition": "“I’ll now apply that operating framework to the longer-duration crop: Aquilaria crassna agarwood.”",
        "sources": [
            f"{DOCS['company']}, pp. 1–2 — Agroforestry Intelligence definition and AI-enabled monitoring differentiator.",
            f"{DOCS['agarwood']}, p. 14 — AI-enabled plantation-intelligence stack.",
            f"{DOCS['mango']}, p. 12 — AI-enabled monitoring for the mango programme.",
            f"{DOCS['presentation']}, p. 26 — technology advantage.",
            f"{DOCS['faq']}, p. 8 — explanation of Agroforestry Intelligence and operations technology.",
        ],
    },
    {
        "group": "Programs",
        "title": "Agarwood Programme — Ten-Year Biological and Commercial Cycle",
        "route": "/agarwood-life-cycle",
        "why": "This page gives a focused explanation of what agarwood is, why Aquilaria crassna is cultivated, how resin is induced and harvested, how later-stage costs are reserved, and what risks qualify the financial illustrations.",
        "inside": [
            "General overview: uses, demand regions, CITES-constrained wild supply, controlled inoculation, and the 23,000-tree launch inventory plus 4,600 replacement trees.",
            "Commercial model: USD 301.71 Year 0 subscription price per tree-equivalent share, including a USD 173.10 segregated reserve for approved later-stage costs.",
            "Illustrative benchmarks: approximately 1.5 kg in Year 9 and 2.0 kg in Year 10, with harvest modelled 50/50 across those years.",
            "Illustrative USD 100,000 case: approximately 331 shares, USD 416,336 net income over ten years, and a 17.8% annualised IRR—expressly not guaranteed.",
            "Risk controls: twenty-percent surplus planting, covered-event replacement, conservative assumptions, and reserve controls.",
            "Four stages: propagation/planting, establishment/growth, controlled inoculation, and harvest/extraction.",
        ],
        "demo": "Lead with the biological cycle. Present the commercial illustration only after the operating stages, and read the non-guarantee qualification aloud.",
        "script": [
            "“Aquilaria crassna produces agarwood when resin forms in response to biological stress. Plantation cultivation uses controlled inoculation to make that process more consistent and traceable while reducing reliance on endangered wild supply.”",
            "“The programme runs across ten years: propagation and planting in Year 0, establishment through Year 6, controlled inoculation in Years 7–8, and harvesting and extraction across Years 9–10.”",
            "“The current illustrative subscription price is 301 dollars and 71 cents per tree-equivalent share, including a segregated 173-dollar-and-10-cent reserve for approved later-stage costs. The displayed USD 100,000 case and its income and IRR figures are management illustrations, not promises. Actual outcomes depend on biological performance, realized export sales, costs, reserve rules, fund controls, and final documents.”",
        ],
        "transition": "“The second crop balances that later realization with an earlier and recurring harvest profile.”",
        "sources": [
            f"{DOCS['agarwood']}, pp. 2–16 and 20–23 — formation, markets, supply, plantation model, economics, technology, impact, risk, and disclaimers.",
            f"{DOCS['presentation']}, pp. 10–15 and 23–27 — commodity, market, lifecycle, USD 100,000 illustration, reserve, structure, and risk controls.",
            f"{DOCS['ppm']}, pp. 2, 5–8, 15–16, and 22–24 — dual-crop role, illustration, economics verification, reserve schedule, and pricing schedule.",
            f"{DOCS['faq']}, pp. 2 and 6–12 — demand, ownership, operations, controls, transfers, and distributions.",
        ],
    },
    {
        "group": "Programs",
        "title": "Mango Programme — Sweet Elena Carabao Orchard Model",
        "route": "/mango-program",
        "why": "This page explains the premium mango variety, high-density orchard model, fruiting and mature-yield assumptions, risk controls, shareholder structure, and the long recurring-income timeline.",
        "inside": [
            "Proprietary Carabao × Elena dwarf mango, selectively grafted for sweetness, quality, and higher-density planting.",
            "Zambales rollout at approximately 416 trees per hectare, with induced flowering and real-time monitoring.",
            "Risk controls: twenty-percent replacement buffer, covered-event replacement, conservative baselines, and projections based on only eighty percent of planted trees producing fruit.",
            "Dedicated ring-fenced mango sub-fund; USD 437.08 Year 0 subscription price per tree-equivalent share.",
            "Programme snapshot: 25-year lifecycle, fruiting from Year 5, approximately 30 kg per tree in Year 5, and approximately 100 kg per tree from Year 10 onward.",
            "Illustrative USD 100,000 case of approximately 229 mango shares, with period and 25-year income/IRR figures qualified as non-guaranteed and subject to commissions, deductions, fund requirements, and final documents.",
        ],
        "demo": "Use the snapshot table to explain time, yield, and share structure. State that figures are illustrative, then point to the footnote before discussing any income or IRR.",
        "script": [
            "“The mango programme uses a Sweet Elena Carabao hybrid with grafting and dwarfing techniques designed for premium fruit quality and a high-density orchard layout.”",
            "“The commercial timeline is different from agarwood. Fruiting is modelled from Year 5 at roughly 30 kilograms per productive tree, increasing toward about 100 kilograms from Year 10 onward, with production continuing through the 25-year orchard model.”",
            "“The current Year 0 subscription price shown is 437 dollars and 8 cents per tree-equivalent share. The USD 100,000 illustration is approximately 229 shares. All income and IRR figures remain illustrative and depend on actual yield, fruit quality, pricing, the twenty-percent harvesting commission, permitted deductions, fund requirements, and definitive documents.”",
        ],
        "transition": "“Clients can inspect the operational setting directly, which leads to the plantation-visit programme.”",
        "sources": [
            f"{DOCS['mango']}, pp. 2–15 and 18–21 — variety, markets, density, investment structure, lifecycle, economics, technology, impact, risk, and disclaimers.",
            f"{DOCS['presentation']}, pp. 16–23 and 27 — mango overview, market, lifecycle, USD 100,000 illustration, deductions, combined summary, and risk controls.",
            f"{DOCS['ppm']}, pp. 2, 5–8, 18, and 24–25 — dual-crop profile, illustration, economics verification, capitalization, and mango pricing/deduction schedules.",
            f"{DOCS['faq']}, pp. 2 and 6–12 — demand, ownership, fees, insurance, replacement stock, transfers, and distributions.",
        ],
    },
    {
        "group": "Programs",
        "title": "Plantation Visit — Operational Visibility and Travel Experience",
        "route": "/plantation-visit",
        "why": "This page turns remote reporting into physical visibility. It combines a structured Zambales plantation visit with premium hospitality and optional onward travel through Clark International Airport.",
        "inside": [
            "Introduction to client visits, dedicated staff support, transportation, and a two-night premium hotel stay.",
            "Featured visit video.",
            "Travel-guide overview of Zambales beaches, Sundowners Resort, Clark airport, and onward island routes.",
            "Clark destinations: Cebu, Coron/Busuanga, Boracay via Caticlan, Bohol-Panglao, El Nido, and Surigao.",
            "Suggested route from Manila through Zambales and Clark.",
            "Destination gallery including San Antonio, Botolan, Clark, Cebu/Oslob, Coron, Boracay, Bohol, El Nido, and Surigao/Siargao.",
        ],
        "demo": "Play only a short segment of the featured video, then show the route and two or three destinations. Keep the focus on client visibility, not on presenting the site as a travel agency.",
        "script": [
            "“The visit programme allows professional shareholders to connect reports and imagery with the actual operating environment in Zambales.”",
            "“A structured visit can include transportation, a management briefing, nursery and plantation walkthroughs, and a two-night premium hotel stay. The page also shows how Clark International Airport can connect the visit with major island destinations.”",
            "“The travel content adds personal value, but its main purpose within this portal is operational transparency: it shows that the programme is accessible and that field inspection can be organised.”",
        ],
        "transition": "“Beyond client visibility, the programme also sets out its environmental and community commitments.”",
        "sources": [
            f"{DOCS['presentation']}, p. 31 — client visibility and reporting access.",
            f"{DOCS['faq']}, pp. 6–8 — operations, reporting, and client-service context.",
        ],
        "source_note": "The detailed routes, destinations, resort, video, and gallery are not taken from the July 2026 PDFs in ‘new files’. They are website travel/media content supported by the separate attached visual guide and gallery assets.",
    },
    {
        "group": "Programs",
        "title": "Positive Impact — Reforestation and Local Participation",
        "route": "/impact",
        "why": "This page summarizes the social and environmental outcomes linked to programme scale without mixing them into the crop-economics pages.",
        "inside": [
            "One-for-one reforestation commitment: one native Philippine tree for each corresponding underlying tree represented in the sub-fund allocation model.",
            "Local employment, training, field-care procedures, and responsible production practices.",
            "Responsible-production standards built around certification pathways, traceability, ecological restoration, and sustainability reporting.",
        ],
        "demo": "Use the two large cards to separate environmental and social outcomes, then finish with the standards statement.",
        "script": [
            "“Commercial planting is linked to a one-for-one native-tree reforestation commitment. For each corresponding underlying tree represented in the allocation model, the programme commits to planting one native Philippine tree.”",
            "“The social side is grounded in local employment and field training in Zambales. Teams are trained in plantation care, monitoring, and responsible production practices.”",
            "“The standards section explains the intended direction: traceable outputs, certification pathways, long-term ecological restoration, and transparent sustainability reporting.”",
        ],
        "transition": "“I’ll now show how professional shareholders are onboarded, informed, and supported.”",
        "sources": [
            f"{DOCS['company']}, p. 1 — one-for-one reforestation differentiator and community-development mission.",
            f"{DOCS['agarwood']}, p. 16 — native-species reforestation and environmental impact.",
            f"{DOCS['mango']}, p. 15 — one-for-one native-species reforestation.",
            f"{DOCS['presentation']}, pp. 28–30 — sustainability, social/economic benefits, and preliminary carbon analysis.",
            f"{DOCS['ppm']}, pp. 30–31 — environmental-impact and carbon-sequestration appendix with qualifications.",
        ],
    },
    {
        "group": "Client & Media",
        "title": "Client Services — Governance, Reporting, and Visits",
        "route": "/services",
        "why": "This page explains what a professional shareholder receives at the service and governance level: regulated onboarding, share records, ring-fenced administration, reporting, document access, and structured field visits.",
        "inside": [
            "Sub-fund shareholding and governance: subscription documents, shareholder register, licensed service-provider oversight, crop-level ring-fencing, and controlled transfer processes.",
            "Professional reporting: periodic fund and plantation reports, audit and field verification, harvest/yield/material-exception reporting, and shareholder communications.",
            "Client visitation programme with premium two-night accommodation and transportation.",
            "Logistics from Clark, Manila, and Subic and a typical nursery/plantation/management itinerary.",
        ],
        "demo": "First cover what the client receives in documents and reporting. Treat visitation as an additional transparency mechanism rather than the main service.",
        "script": [
            "“Client services begin with regulated onboarding and fund governance. Eligible professional clients subscribe through licensed private-placement intermediaries and receive shares in the relevant ring-fenced VCC sub-fund.”",
            "“The service model includes evidence of share issuance, entry in the shareholder register, licensed manager and administrator oversight, periodic fund and plantation reports, independent audit or field-verification outputs, and material-exception reporting.”",
            "“The structured visitation programme complements those records by giving clients a practical way to inspect the nursery and plantation and receive a management briefing.”",
        ],
        "transition": "“For direct questions, partnership enquiries, or a registration request, the next page provides the contact workflow.”",
        "sources": [
            f"{DOCS['presentation']}, pp. 24–25 and 31–35 — structure, professional operations, visibility, next steps, and contact.",
            f"{DOCS['ppm']}, pp. 2–3, 11–21, and 26–29 — investor eligibility, structure, reporting, restrictions, risks, and readiness conditions.",
            f"{DOCS['faq']}, pp. 6–12 — ownership, operations, documentation, transfers, subscription process, and distributions.",
        ],
    },
    {
        "group": "Client & Media",
        "title": "Contact Us — Enquiry and Office Information",
        "route": "/contact",
        "why": "This page gives visitors a controlled point of contact for partnerships, operations, client support, and registration requests while clearly identifying the group’s UAE and Philippine offices.",
        "inside": [
            "Registration-request form for first name, last name, email, and message, with success and error feedback.",
            "Workflow statement that the team reviews the request and responds through the appropriate Pipedrive process.",
            "Telephone numbers, office email, and website address.",
            "Golden Forests ADGM holding-company address in Abu Dhabi.",
            "Sales and marketing office in Dubai Digital Park.",
            "CADI plantation-management office at the Philippine Stock Exchange Tower in BGC, Taguig.",
        ],
        "demo": "Do not submit a test enquiry during a live presentation. Show the form fields, then scroll to the office identities and contact details.",
        "script": [
            "“The contact page is the formal entry point for partnership, operational, and client-support enquiries. A visitor can submit a registration request, which is reviewed and routed through the appropriate workflow.”",
            "“The page also separates the group’s principal entities and locations: the ADGM holding company in Abu Dhabi, the sales and marketing operation in Dubai, and CADI’s plantation-management office in BGC, Taguig.”",
            "“This gives the audience both a digital enquiry route and clearly identified office contact points.”",
        ],
        "transition": "“For visual evidence of ongoing work, I’ll move next to the operational photo gallery.”",
        "sources": [
            f"{DOCS['presentation']}, p. 35 — next steps and contact context.",
            f"{DOCS['company']}, pp. 1–2 — Golden Forests’ Singapore, Dubai, and Manila operating footprint.",
        ],
        "source_note": "Exact telephone numbers, email addresses, workflow wording, and street addresses are maintained in the website copy and are not all reproduced in the July 2026 PDFs.",
    },
    {
        "group": "Client & Media",
        "title": "Operational Photo Gallery — Visual Evidence Library",
        "route": "/photo-gallery",
        "why": "This page centralizes visual proof of propagation, field rollout, facilities, and personnel. It allows operational media to be searched and grouped without crowding the narrative pages.",
        "inside": [
            "Collection statistics for total media, categories, videos, and the current collection.",
            "Four browsable collections: Nursery, Plantation, Facilities, and Team.",
            "Search by title, description, or location.",
            "Media cards with full-view opening behavior and empty-state messaging when no results match.",
            "Admin-managed ordering and asset library so the gallery can be updated without code changes.",
        ],
        "demo": "Switch between Nursery and Plantation, search for one location or subject, and open a single image in the full media view.",
        "script": [
            "“The gallery is the visual archive for the operating programme. It separates nursery propagation, plantation work, facilities, and team documentation into searchable collections.”",
            "“This is useful because the operational narrative can be checked against current photographs and videos without overloading the main pages.”",
            "“The media library is maintained through the private administration area, so new field evidence and updated ordering can be published as the programme progresses.”",
        ],
        "transition": "“I’ll close the public walkthrough with the risk warning and disclaimer that applies across the portal.”",
        "sources": [
            f"{DOCS['presentation']}, pp. 25, 26, and 31 — professional operations, technology evidence, and client visibility.",
            f"{DOCS['ppm']}, pp. 9 and 17 — current operating base and contemplated investor reporting/communications.",
        ],
        "source_note": "The individual photos, videos, titles, categories, and counts come from the website gallery assets and database, not from the PDFs in ‘new files’.",
    },
    {
        "group": "Legal",
        "title": "Risk Warning and Disclaimer — Boundaries of the Website",
        "route": "/disclaimer",
        "why": "This page makes clear that the portal is for operational information and transparency, not a public offer, financial promotion, investment recommendation, or guarantee of performance.",
        "inside": [
            "Important notice identifying CADI as the Philippine operational platform.",
            "Warnings concerning projections, biological assets, changing operational data, and non-guaranteed outcomes.",
            "Nature of investment: pooled ring-fenced sub-fund shares, no direct ownership of trees or land, long-term illiquidity, and material biological, environmental, market, operational, counterparty, and regulatory risks.",
            "Regulatory-compliance statement, professional-investor/jurisdiction restrictions, confidentiality, no-liability language, and contact email.",
            "A shorter version of the warning also appears in the site-wide footer on every public page.",
        ],
        "demo": "Show the numbered sections and emphasize the first, second, fourth, and sixth sections. Do not paraphrase the legal wording if a precise disclosure question is asked.",
        "script": [
            "“The website is published by CADI for operational information and transparency. It is not a public offer, financial promotion, investment advice, or a solicitation to invest.”",
            "“Any yield, income, IRR, market, or operational figure is based on assumptions and can differ materially from actual results. Sub-fund shares do not convey ownership of a particular tree or land and are long-term, illiquid, and exposed to biological, environmental, market, operational, counterparty, and regulatory risks.”",
            "“Any potential subscription is limited to eligible professional investors, through appropriately licensed channels, and only on the basis of final offering and transaction documents. The website should therefore be read as an operational portal, with the final legal documents controlling.”",
        ],
        "transition": "“That completes the public website walkthrough. I’ll now summarize the full structure and the private operational tools behind it.”",
        "sources": [
            f"{DOCS['company']}, p. 2 — disclosure and risk-factor notice.",
            f"{DOCS['agarwood']}, pp. 22–23 — notices, forward-looking statements, and disclaimers.",
            f"{DOCS['mango']}, pp. 20–21 — important notices and disclaimers.",
            f"{DOCS['presentation']}, pp. 34–35 — next steps, professional-client context, and non-guarantee notice.",
            f"{DOCS['ppm']}, pp. 1, 3, and 18–21 — draft status, distribution limits, selling restrictions, and risk factors.",
        ],
    },
]


def configure_document(doc):
    sec = doc.sections[0]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.7)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(0.8)
    sec.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(CHARCOAL)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color in (
        ("Title", 30, GREEN),
        ("Heading 1", 23, GREEN),
        ("Heading 2", 14, GREEN),
        ("Heading 3", 11.5, GREEN),
    ):
        s = styles[style_name]
        s.font.name = "Aptos Display"
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.font.bold = True
        s.paragraph_format.space_before = Pt(10)
        s.paragraph_format.space_after = Pt(5)
        s.paragraph_format.keep_with_next = True

    styles["Title"].font.name = "Georgia"
    styles["Title"].font.bold = True

    custom = [
        ("Eyebrow", 9, GOLD, True),
        ("Route", 9.5, MUTED, False),
        ("Script Label", 9.5, GREEN, True),
        ("Script Text", 10.5, CHARCOAL, False),
        ("Source Text", 8.5, MUTED, False),
        ("TOC Entry", 10.5, CHARCOAL, False),
    ]
    for name, size, color, bold in custom:
        if name not in styles:
            s = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            s = styles[name]
        s.font.name = "Aptos"
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.font.bold = bold
        s.paragraph_format.space_after = Pt(4)
    styles["Eyebrow"].font.letter_spacing = Pt(1.2)
    styles["Script Text"].paragraph_format.line_spacing = 1.14
    styles["Script Text"].paragraph_format.space_after = Pt(7)

    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("GOLDEN FORESTS  |  WEBSITE PRESENTATION SCRIPT")
        r.font.name = "Aptos"
        r.font.size = Pt(7.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(GOLD)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = fp.add_run("Crassna Agroforestry Development Inc.  •  Internal presentation guide  •  ")
        rr.font.name = "Aptos"
        rr.font.size = Pt(7.5)
        rr.font.color.rgb = RGBColor.from_string(MUTED)
        add_field(fp, "PAGE")


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(45)
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(1.05))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run("GOLDEN FORESTS")
    r.font.name = "Aptos"
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(GOLD)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Whole Website Structure\nand Presentation Script")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(8)
    run = sub.add_run("A page-by-page speaking guide for the CADI Plantation Management Portal")
    run.font.name = "Aptos"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = 1
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(5.7)
    shade(cell, GREEN)
    set_cell_margins(cell, 220, 260, 220, 260)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared from the live website code and the July 2026 documents in the “new files” folder")
    r.font.name = "Aptos"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(PALE_GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(100)
    r = p.add_run("Prepared 21 July 2026  |  Asia/Manila")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(MUTED)


def add_how_to_use(doc):
    doc.add_page_break()
    doc.add_heading("How to Use This Script", level=1)
    doc.add_paragraph(
        "This document is written for a live website presentation. Follow the public sidebar from top to bottom. "
        "For every page, the guide identifies the purpose, visible content, demonstration cue, suggested speaking script, transition, and source lineage."
    )
    add_bullets(doc, [
        "Words inside quotation marks are the recommended talk track. Adapt the tone, but preserve legal qualifications.",
        "Check every live metric and its update date immediately before presenting; database values can change after this document is generated.",
        "When stating projected income, IRR, yield, price, or timing, always include the accompanying non-guarantee language.",
        "Say “tree-equivalent share” or “a share valued by reference to an underlying tree.” Do not say that a shareholder owns an individual tree, land, or planting block.",
        "The draft PPM is a discussion draft. It is not a final offer, subscription document, or permission to market publicly.",
        "Detailed citations refer to PDF page order as stored in the files, which may differ from a printed page number shown inside a PDF.",
    ])

    doc.add_heading("Recommended presentation flow", level=2)
    flow = doc.add_table(rows=1, cols=3)
    flow.style = "Table Grid"
    headers = ["Phase", "Pages", "Purpose"]
    for i, text in enumerate(headers):
        shade(flow.rows[0].cells[i], GREEN)
        p = flow.rows[0].cells[i].paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(flow.rows[0])
    rows = [
        ("1. Orient", "Home, About, Management", "Who we are, how the structure works, and who executes."),
        ("2. Evidence", "Nursery, Plantation, Compliance, Technology", "What is operating, how it is controlled, and how it is measured."),
        ("3. Programmes", "Agarwood, Mango, Plantation Visit, Impact", "Crop timelines, qualified economics, field access, and outcomes."),
        ("4. Support", "Client Services, Contact, Gallery", "Shareholder administration, enquiries, and visual evidence."),
        ("5. Close", "Disclaimer", "Risk boundaries and controlling-document reminder."),
    ]
    for row in rows:
        cells = flow.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            set_cell_margins(cells[i])

    doc.add_heading("Opening script", level=2)
    add_script_box(doc, [
        "“Good [morning/afternoon]. Today I’ll walk you through the Golden Forests plantation management website operated by Crassna Agroforestry Development Inc., or CADI, in the Philippines.”",
        "“The portal is structured to move from the corporate overview into operational evidence, crop programmes, client services, and risk disclosures. It is an operational-information website for eligible professional-investor context; it is not a public offer or a guarantee of returns.”",
    ])


def add_structure_map(doc):
    doc.add_page_break()
    doc.add_heading("Complete Public Website Map", level=1)
    doc.add_paragraph("The persistent left sidebar groups the public presentation into four sections. A common header, footer, responsive sidebar, and disclaimer link appear throughout the public site.")
    groups = {}
    for p in pages:
        groups.setdefault(p["group"], []).append(p)
    for group in ("Overview", "Operations", "Programs", "Client & Media", "Legal"):
        doc.add_heading(group, level=2)
        for p in groups.get(group, []):
            x = doc.add_paragraph(style="TOC Entry")
            x.add_run(f"{p['route']}  —  ").bold = True
            x.add_run(p["title"].split(" — ")[0])

    doc.add_heading("Shared site-wide elements", level=2)
    add_bullets(doc, [
        "Header: sidebar toggle, Golden Forests logo, ‘Plantation Management Portal,’ CADI company name, and ‘Powered by Golden Forests.’",
        "Sidebar: brand panel plus grouped navigation for Overview, Operations, Programs, and Client & Media.",
        "Footer: CADI/Golden Forests identity, links to About, Contact, Gallery, and Disclaimer, plus the short risk notice.",
        "Responsive behavior: collapsible sidebar and mobile navigation.",
        "Real-time content synchronization: public pages refresh when managed content changes.",
    ])


def add_appendices(doc):
    doc.add_page_break()
    doc.add_heading("Appendix A — Private Administration Area", level=1)
    doc.add_paragraph("The following routes are part of the complete website structure but should not be opened during a public presentation unless the audience is specifically authorised.")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, text in enumerate(("Route / area", "What is inside", "Why it exists")):
        shade(table.rows[0].cells[i], GREEN)
        p = table.rows[0].cells[i].paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("/admin/login", "Username/password login and authentication feedback.", "Restricts content management to authorised personnel."),
        ("Nursery Metrics", "Agarwood height, mango height, mortality rate, stock counts, sales/share availability, and update dates.", "Keeps the live nursery dashboard current."),
        ("Operational Updates", "Create, edit, and list dated updates by category.", "Supports an operational log and future transparency features."),
        ("Team Profiles", "Create/edit team names, titles, categories, biographies, expertise, images, and ordering.", "Maintains the public Management page without code changes."),
        ("Gallery Manager", "Upload, categorize, describe, order, and remove media; control the live seedling gallery.", "Maintains operational visual evidence."),
        ("Contact Inbox", "Review and delete contact submissions.", "Centralizes enquiries received from the public Contact page."),
        ("Website Copy", "Edit page copy and typography settings.", "Allows controlled updates to public content while preserving the page structure."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            set_cell_margins(cells[i])

    doc.add_heading("Admin presentation note", level=2)
    add_script_box(doc, [
        "“Behind the public portal is a private administration console. It allows authorised personnel to update nursery metrics, operational updates, team profiles, gallery media, contact submissions, and website copy without rebuilding the site.”",
        "“Because the console contains operational controls and contact data, it is separated from the public layout and should only be demonstrated in an authorised session.”",
    ])

    doc.add_page_break()
    doc.add_heading("Appendix B — Routes, Aliases, and Current Implementation Notes", level=1)
    notes = [
        ("/location", "Currently resolves to the same Plantation Visit/Ecotourism page as /plantation-visit."),
        ("/ecotourism", "Also resolves to the same Plantation Visit/Ecotourism page."),
        ("Location.tsx", "A separate legacy ‘Zambales & Key Island Destinations’ page exists in the source but is not connected to the active router."),
        ("OperationalUpdatesLog.tsx", "An operational-updates page exists in the source but has no active public route."),
        ("/ai-technology", "The Home operations-snapshot technology card currently points here, but the active Technology route is /technology. Clicking the card can therefore reach the 404 page until corrected."),
        ("404 fallback", "Any unmatched route displays ‘404 Page Not Found’ with a link back to Home."),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    for i, text in enumerate(("Item", "Current behavior")):
        shade(t.rows[0].cells[i], GREEN)
        p = t.rows[0].cells[i].paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(t.rows[0])
    for item, behavior in notes:
        cells = t.add_row().cells
        cells[0].text = item
        cells[1].text = behavior
        set_cell_margins(cells[0])
        set_cell_margins(cells[1])
    p = doc.add_paragraph()
    r = p.add_run("Presenter action before going live: ")
    r.bold = True
    p.add_run("Use the sidebar’s /technology link rather than the Home technology card unless the route is corrected. Avoid demonstrating legacy or unlinked pages as if they are part of the public navigation.")

    doc.add_page_break()
    doc.add_heading("Appendix C — Source Register", level=1)
    doc.add_paragraph("Authoritative July 2026 source set found in the project’s “new files” folder:")
    source_descriptions = [
        (DOCS["company"], "Concise corporate overview, vision, mission, differentiators, assets/share classes, terminology, and disclosure."),
        (DOCS["agarwood"], "Agarwood formation, market, supply, plantation model, economics, technology, impact, risk, and disclaimers."),
        (DOCS["mango"], "Sweet Elena variety, market, orchard model, economics, technology, impact, risk, and disclaimers."),
        (DOCS["presentation"], "Professional-client presentation covering group, two-crop strategy, financial illustrations, structure, operations, technology, impact, visibility, approvals, and next steps."),
        (DOCS["faq"], "Professional/corporate-investor questions covering company, ownership, operations, technology, fees, sustainability, documentation, transfer, subscription, and distributions."),
        (DOCS["ppm"], "Draft proposed VCC strategy, economics, structure, controls, reporting, restrictions, risks, pricing schedules, readiness conditions, and environmental appendix."),
    ]
    for name, description in source_descriptions:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(name)
        r.bold = True
        p.add_run(" — " + description)

    doc.add_heading("Source-control cautions", level=2)
    add_bullets(doc, [
        "The draft PPM repeatedly states that it is subject to legal, regulatory, tax, structuring, and commercial review and is not a final offer document.",
        "Some exposé wording uses legacy direct-tree-ownership language. The live website intentionally uses the updated shareholding model: shareholders hold sub-fund shares and do not directly own individual trees.",
        "Where figures differ among documents, this script follows the current live website copy and the later professional-client/PPM figures, while preserving non-guarantee language.",
        "The current website shows the mango subscription price as USD 437.08. The two-page company overview contains a different mango figure (USD 371.64); it should be reconciled before external reuse of that overview.",
        "Certification dates and regulatory statements should be re-verified against current evidence before an external presentation.",
    ])

    doc.add_page_break()
    doc.add_heading("Appendix D — Presenter’s Final Checklist", level=1)
    checklist = [
        "Confirm the website is online and all public routes load.",
        "Record the current nursery-metric update date and current live values.",
        "Confirm the correct mango share price and other economics against approved final materials.",
        "Check that management profiles and job titles are current.",
        "Test the featured videos and gallery media on the presentation device.",
        "Use /technology from the sidebar; verify whether the /ai-technology Home-card link has been fixed.",
        "Do not expose the admin console, contact inbox, credentials, or confidential commercial data.",
        "Keep the Risk Warning and Disclaimer available for questions.",
        "Do not describe projected income, yield, IRR, harvest timing, certification, or regulatory outcomes as guaranteed.",
        "Close by directing eligible parties to licensed channels and final offering documents, not to the website as an offer document.",
    ]
    for item in checklist:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        p.add_run("☐  ").font.color.rgb = RGBColor.from_string(GOLD)
        p.add_run(item)

    doc.add_heading("Closing script", level=2)
    add_script_box(doc, [
        "“In summary, the website is designed to connect the Golden Forests structure with visible Philippine operations: who manages the programme, how planting stock and field execution are tracked, how each crop develops, how clients receive information, and which risks and legal boundaries apply.”",
        "“For any potential transaction, the controlling materials are the final approved offering, subscription, and governance documents, together with the required professional-investor eligibility and licensed-placement process. Thank you, and I’m happy to take questions.”",
    ])


def set_core_properties(doc):
    props = doc.core_properties
    props.title = "Golden Forests Whole Website Structure and Presentation Script"
    props.subject = "Page-by-page presentation guide with source lineage"
    props.author = "Crassna Agroforestry Development Inc."
    props.keywords = "Golden Forests, CADI, website, presentation, structure, script, July 2026"
    props.comments = "Generated from the live project structure and July 2026 source documents."


def main():
    doc = Document()
    configure_document(doc)
    set_core_properties(doc)
    add_cover(doc)
    add_how_to_use(doc)
    add_structure_map(doc)
    for i, page in enumerate(pages, 1):
        add_page_section(doc, i, page)
    add_appendices(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
